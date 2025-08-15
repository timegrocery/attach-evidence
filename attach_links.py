import re, sys, traceback
from pathlib import Path
from openpyxl import load_workbook
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.run import Run

import ttkbootstrap as tb
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
from tkinter.filedialog import askopenfilename, asksaveasfilename
from tkinter.scrolledtext import ScrolledText

# ---------------- Core logic ----------------
PATTERN = re.compile(r"\[(?P<code>\d+\.\d+-\d{3})\]")  # [1.2-001]

def read_excel_headers(excel_path: str):
    """Đọc hàng tiêu đề (row 1) bằng openpyxl, trả về list header strings."""
    wb = load_workbook(excel_path, read_only=True, data_only=True)
    ws = wb.active
    row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True), None)
    headers = [str(v).strip() if v is not None else "" for v in (row or [])]
    wb.close()
    return headers

def load_mapping(excel_path, code_col, link_col, summary_col=None):
    wb = load_workbook(excel_path, read_only=True, data_only=True)
    ws = wb.active

    headers = [str(c.value).strip() if c.value is not None else "" for c in next(ws.iter_rows(min_row=1, max_row=1))]
    idx = {h.lower(): i for i, h in enumerate(headers)}

    def need(col):
        k = str(col).strip().lower()
        if k not in idx:
            raise ValueError(f"Không tìm thấy cột '{col}' trong Excel.")
        return idx[k]

    i_code = need(code_col)
    i_link = need(link_col)
    i_sum  = idx.get(str(summary_col).strip().lower()) if summary_col else None

    mapping = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        code = (str(row[i_code]).strip("[]").strip() if row[i_code] is not None else "")
        link = (str(row[i_link]).strip()           if row[i_link] is not None else "")
        if not code or not link:
            continue
        tip  = (str(row[i_sum]).strip() if (i_sum is not None and row[i_sum] is not None) else None)
        mapping[code] = {"url": link, "tip": tip}
    wb.close()
    return mapping, headers

def add_hyperlink(paragraph, anchor_run: Run, url: str, text: str, tooltip: str = None):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    if tooltip:
        hyperlink.set(qn('w:tooltip'), tooltip)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0000FF')
    underline = OxmlElement('w:u'); underline.set(qn('w:val'), 'single')
    rPr.append(color); rPr.append(underline)
    t = OxmlElement('w:t'); t.text = text
    new_run.append(rPr); new_run.append(t)
    hyperlink.append(new_run)

    p = paragraph._p
    idx = p.index(anchor_run._r)
    p.insert(idx, hyperlink)
    p.remove(anchor_run._r)

def replace_paragraph_with_hyperlinks(paragraph, mapping):
    text = "".join(run.text for run in paragraph.runs)
    if not text: return 0, 0
    matches = list(PATTERN.finditer(text))
    if not matches: return 0, 0
    base_font = paragraph.runs[0].font if paragraph.runs else None

    for r in paragraph.runs: r.text = ""
    paragraph.add_run("")

    found = linked = 0
    last = 0
    for m in matches:
        before = text[last:m.start()]
        if before:
            r = paragraph.add_run(before)
            if base_font: r.font.name, r.font.size = base_font.name, base_font.size

        code_key = m.group("code")
        token = m.group(0)
        found += 1

        info = mapping.get(code_key) or mapping.get(code_key.strip())
        if info:
            url = info["url"] if isinstance(info, dict) else str(info)
            tip = info.get("tip") if isinstance(info, dict) else None
            if url:
                anchor = paragraph.add_run("")
                add_hyperlink(paragraph, anchor, url, token, tooltip=tip)
                linked += 1
            else:
                r = paragraph.add_run(token)
                if base_font: r.font.name, r.font.size = base_font.name, base_font.size
        else:
            r = paragraph.add_run(token)
            if base_font: r.font.name, r.font.size = base_font.name, base_font.size

        last = m.end()

    tail = text[last:]
    if tail:
        r = paragraph.add_run(tail)
        if base_font: r.font.name, r.font.size = base_font.name, base_font.size

    return found, linked

def process_document(doc_path, out_path, mapping, log=None):
    def logprint(*a):
        if log is not None:
            log.configure(state='normal'); log.insert('end', " ".join(map(str,a)) + "\n")
            log.see('end'); log.configure(state='disabled')
        else:
            print(*a)

    doc = Document(doc_path)
    total_found = total_linked = 0

    for p in doc.paragraphs:
        f, l = replace_paragraph_with_hyperlinks(p, mapping)
        total_found += f; total_linked += l

    for i, table in enumerate(doc.tables, 1):
        logprint(f"Đang xử lý bảng {i} ...")
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    f, l = replace_paragraph_with_hyperlinks(p, mapping)
                    total_found += f; total_linked += l

    try:
        doc.save(out_path)
    except PermissionError:
        raise PermissionError(
            "Không thể ghi file đầu ra. Có thể file đang được mở trong Word.\n"
            "Vui lòng ĐÓNG file Word rồi chạy lại."
        )

    return total_found, total_linked

# ---------------- UI ----------------
class App(tb.Window):
    def __init__(self):
        super().__init__(themename="journal")
        self.title("Attach Evidence Links")
        self.minsize(980, 560)

        header = tb.Label(self, text="Attach Evidence Links", font=("Segoe UI", 16, "bold"))
        header.pack(side=TOP, anchor='w', padx=16, pady=(12,6))

        body = tb.Frame(self, padding=10)
        body.pack(fill=BOTH, expand=YES)

        form = tb.Labelframe(body, text="Nguồn dữ liệu", padding=12, bootstyle=INFO)
        form.pack(fill=X)

        # Word
        self.word_var = tb.StringVar()
        tb.Label(form, text="Word (.docx):").grid(row=0, column=0, sticky='e', padx=6, pady=6)
        tb.Entry(form, textvariable=self.word_var).grid(row=0, column=1, columnspan=3, sticky='ew', padx=6, pady=6)
        tb.Button(form, text="Browse", command=self.pick_word, bootstyle=SECONDARY).grid(row=0, column=4, padx=6, pady=6)

        # Excel
        self.excel_var = tb.StringVar()
        tb.Label(form, text="Excel (.xlsx):").grid(row=1, column=0, sticky='e', padx=6, pady=6)
        tb.Entry(form, textvariable=self.excel_var).grid(row=1, column=1, columnspan=3, sticky='ew', padx=6, pady=6)
        tb.Button(form, text="Browse", command=self.pick_excel, bootstyle=SECONDARY).grid(row=1, column=4, padx=6, pady=6)

        # Columns
        tb.Label(form, text="Cột mã (*):").grid(row=2, column=0, sticky='e', padx=6, pady=6)
        self.col_code = tb.Combobox(form, state='readonly', width=30)
        self.col_code.grid(row=2, column=1, sticky='w', padx=6, pady=6)

        tb.Label(form, text="Cột link (*):").grid(row=2, column=2, sticky='e', padx=6, pady=6)
        self.col_link = tb.Combobox(form, state='readonly', width=30)
        self.col_link.grid(row=2, column=3, sticky='w', padx=6, pady=6)

        tb.Label(form, text="Cột mô tả (Tùy chọn):").grid(row=3, column=0, sticky='e', padx=6, pady=6)
        self.col_sum = tb.Combobox(form, state='readonly', width=30)
        self.col_sum.grid(row=3, column=1, sticky='w', padx=6, pady=6)

        # Output
        self.out_var = tb.StringVar()
        tb.Label(form, text="File đầu ra:").grid(row=4, column=0, sticky='e', padx=6, pady=6)
        tb.Entry(form, textvariable=self.out_var).grid(row=4, column=1, columnspan=3, sticky='ew', padx=6, pady=6)
        tb.Button(form, text="Save As...", command=self.pick_output, bootstyle=SECONDARY).grid(row=4, column=4, padx=6, pady=6)

        for c in (1,2,3):
            form.grid_columnconfigure(c, weight=1)

        # Actions
        actions = tb.Frame(body, padding=(0,10,0,6))
        actions.pack(fill=X)
        self.run_btn = tb.Button(actions, text="Run", command=self.run, bootstyle=SUCCESS)
        self.run_btn.pack(side=LEFT, padx=(0,8))
        tb.Button(actions, text="Thoát", command=self.destroy, bootstyle=SECONDARY).pack(side=LEFT)
        self.prog = tb.Progressbar(actions, mode="indeterminate", bootstyle=SUCCESS, length=220)
        self.prog.pack(side=RIGHT)

        # Log
        log_frame = tb.Labelframe(body, text="Log", padding=10, bootstyle=INFO)
        log_frame.pack(fill=BOTH, expand=YES)
        self.log = ScrolledText(log_frame, height=14, font=("Consolas", 10))
        self.log.pack(fill=BOTH, expand=YES)

        tb.Label(self, text="Tip: rê chuột lên mã đã gắn link để xem ScreenTip.", foreground="#666")\
          .pack(side=BOTTOM, anchor='w', padx=16, pady=(0,10))

    # ------- helpers -------
    def logprint(self, *a):
        self.log.configure(state='normal')
        self.log.insert('end', " ".join(map(str,a)) + "\n")
        self.log.see('end')
        self.log.configure(state='disabled')

    def pick_word(self):
        p = askopenfilename(title="Chọn file Word", filetypes=[("Word", "*.docx")])
        if p:
            self.word_var.set(p)
            if not self.out_var.get():
                self.out_var.set(str(Path(p).with_name(Path(p).stem + "_linked.docx")))

    def pick_excel(self):
        p = askopenfilename(title="Chọn file Excel", filetypes=[("Excel", "*.xlsx")])
        if not p:
            return
        self.excel_var.set(p)
        try:
            cols = read_excel_headers(p)  # <-- dùng openpyxl
            self.col_code.configure(values=cols)
            self.col_link.configure(values=cols)
            self.col_sum.configure(values=[''] + cols)

            def pick(keys):
                for g in keys:
                    for c in cols:
                        if g.lower() in str(c).lower():
                            return c
                return ''
            self.col_code.set(pick(['mã','code','evidence code']) or (cols[0] if cols else ''))
            self.col_link.set(pick(['link','url']) or (cols[1] if len(cols)>1 else (cols[0] if cols else '')))
            self.col_sum.set(pick(['evidence summary','summary','mô tả','mo ta']) or '')

            self.logprint("Đã nạp tiêu đề cột từ Excel.")
        except Exception as e:
            Messagebox.show_error(f"Không thể đọc tiêu đề cột từ Excel:\n{e}", "Lỗi")

    def pick_output(self):
        p = asksaveasfilename(title="Chọn nơi lưu", defaultextension=".docx",
                              filetypes=[("Word", "*.docx")])
        if p:
            self.out_var.set(p)

    def run(self):
        try:
            word_path = self.word_var.get().strip()
            excel_path = self.excel_var.get().strip()
            col_code   = self.col_code.get().strip()
            col_link   = self.col_link.get().strip()
            col_sum    = self.col_sum.get().strip() or None
            out_path   = self.out_var.get().strip()

            if not word_path or not Path(word_path).is_file():
                Messagebox.show_error("Hãy chọn file Word (.docx) hợp lệ.", "Lỗi"); return
            if not excel_path or not Path(excel_path).is_file():
                Messagebox.show_error("Hãy chọn file Excel (.xlsx) hợp lệ.", "Lỗi"); return
            if not col_code or not col_link:
                Messagebox.show_error("Hãy chọn tên cột mã và cột link.", "Lỗi"); return
            if not out_path:
                out_path = str(Path(word_path).with_name(Path(word_path).stem + "_linked.docx"))
                self.out_var.set(out_path)

            self.run_btn.configure(state=DISABLED)
            self.prog.start(12)

            self.logprint("Đang đọc Excel và tạo mapping...")
            mapping, _ = load_mapping(excel_path, col_code, col_link, col_sum)
            self.logprint(f"Đã nạp {len(mapping)} mã từ Excel.")
            self.logprint("Đang xử lý Word, vui lòng chờ...")

            try:
                found, linked = process_document(word_path, out_path, mapping, log=self.log)
            except PermissionError as pe:
                Messagebox.show_error(f"{pe}\n\nFile đang bị khóa:\n{out_path}", "Không thể ghi file")
                return

            self.logprint(f"Hoàn tất! Tìm thấy {found} mã; đã gắn liên kết cho {linked} mã.")
            if found != linked:
                self.logprint("Lưu ý: một số mã không tìm thấy link trong Excel nên giữ nguyên dạng text.")
            self.logprint(f"File đã lưu: {out_path}")
            Messagebox.show_info("Hoàn tất.", "Xong!")
        except Exception:
            err = ''.join(traceback.format_exception(*sys.exc_info()))
            self.logprint("ĐÃ XẢY RA LỖI:\n", err)
            Messagebox.show_error("Có lỗi xảy ra. Vui lòng xem chi tiết trong ô log.", "Lỗi")
        finally:
            self.prog.stop()
            self.run_btn.configure(state=NORMAL)

if __name__ == "__main__":
    App().mainloop()
